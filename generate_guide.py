import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors

# ================= 内容配置 =================
TITLE = "“蒸蒸日上的胜率” 机器人使用指南"
SECTIONS = [
    {
        "title": "1. 📝 记录战绩 (#add)",
        "content": "一局游戏结束后，请直接复制粘贴玩家信息。\n\n• 格式：玩家名 主将 副将 [胜者标记]\n• 胜者标记：在获胜者的行尾emoji, 或 win 均可识别。\n• 智能识别：机器人能自动识别武将昵称（如“大宝”、“香香”、“双头”）。",
        "code": "#add\n殊 孙尚香 丁奉💅\nMU 刘备 魏延\nDH 狗货 刚烈\nJX 黄忠 祝融\nZS 邹氏 孔融\nJoyce 大宝 凌统💅\nWu 卧龙 庞统\nJason 司马 邓艾"
    },
    {
        "title": "2. 📅 查看今日战报 (#date)",
        "content": "生成一张包含今日所有对局的 Excel 风格图片，直观展示每个人的武将组合和胜负情况。\n\n• 查看今天：发送 #date\n• 查看指定日期：发送 #date 2025-12-09",
        "code": ""
    },
    {
        "title": "3. 🔍 查看单局详情 (#game)",
        "content": "如果你想回看某一局的具体配置。\n\n• 查看最新一局：发送 #game last\n• 查看指定局号：发送 #game 5",
        "code": ""
    },
    {
        "title": "4. 📊 查看胜率排行榜 (#rates)",
        "content": "展示所有玩家的“胜场/总场”以及“胜率”排名。",
        "code": "#rates"
    },
    {
        "title": "5. 📧 导出 Excel 到邮箱 (#email)",
        "content": "将所有历史对局记录生成 Excel 表格，发送到你的邮箱。表格中获胜的武将会高亮显示。",
        "code": "#email yourname@gmail.com"
    },
    {
        "title": "6. ✏️ 修改与删除 (#update / #remove)",
        "content": "如果记录记错了，可以使用以下指令进行修正。机器人会弹出确认按钮。\n\n• 修改某局信息（先输入指令，换行后输入新数据）：\n#update 10\n殊 孙尚香 丁奉\nMU 刘备 赵云 💅\n...\n\n• 删除某局信息：\n#remove 10",
        "code": ""
    },
    {
        "title": "💡 小贴士",
        "content": "1. 武将昵称：你可以直接用黑话，比如“香香”会自动变成“孙尚香”. \n2. 自动时间：录入时机器人会自动记录当前时间，不需要手动输入时间戳。",
        "code": ""
    }
]

# ================= 生成 Word (.docx) =================
def create_word():
    doc = Document()
    
    # 标题
    heading = doc.add_heading(TITLE, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("欢迎使用国战胜率记录助手！以下是所有可用指令的说明。\n(注：所有指令均以 # 开头，支持中文或拼音输入)")

    for sec in SECTIONS:
        doc.add_heading(sec["title"], level=1)
        doc.add_paragraph(sec["content"])
        if sec["code"]:
            # 模拟代码块格式
            p = doc.add_paragraph()
            runner = p.add_run(sec["code"])
            runner.font.name = 'Courier New'
            runner.font.size = Pt(9)
            runner.font.color.rgb = RGBColor(0, 100, 0) # 深绿色代码
            p.paragraph_format.left_indent = Pt(20) # 缩进

    filename = "Sanguosha_Bot_User_Guide.docx"
    doc.save(filename)
    print(f"✅ Word 文档已生成: {filename}")

# ================= 生成 PDF (.pdf) =================
def create_pdf():
    filename = "Sanguosha_Bot_User_Guide.pdf"
    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4
    
    # 注册中文字体 (需要 simhei.ttf 在同级目录，否则 PDF 中文会乱码)
    # 如果没有字体文件，这步会报错。建议优先使用 Word 生成。
    try:
        pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
        font_name = 'SimHei'
    except:
        print("⚠️ 未找到 simhei.ttf，PDF 中文可能无法显示。尝试使用默认字体...")
        font_name = 'Helvetica' # 不支持中文

    y = height - 50
    
    # 标题
    c.setFont(font_name, 18)
    c.drawCentredString(width / 2, y, TITLE)
    y -= 40
    
    c.setFont(font_name, 10)
    c.drawString(50, y, "欢迎使用国战胜率记录助手！以下是所有可用指令的说明。")
    y -= 30

    for sec in SECTIONS:
        if y < 100: # 换页
            c.showPage()
            y = height - 50
            c.setFont(font_name, 10)

        # 章节标题
        c.setFont(font_name, 14)
        c.setFillColor(colors.darkblue)
        c.drawString(50, y, sec["title"])
        y -= 20
        
        # 内容 (简单的换行处理)
        c.setFont(font_name, 10)
        c.setFillColor(colors.black)
        lines = sec["content"].split('\n')
        for line in lines:
            c.drawString(50, y, line)
            y -= 15
            
        # 代码块
        if sec["code"]:
            y -= 5
            c.setFillColor(colors.darkgreen)
            code_lines = sec["code"].split('\n')
            for cl in code_lines:
                c.drawString(70, y, cl)
                y -= 12
        
        y -= 20 # 段落间距

    c.save()
    print(f"✅ PDF 文档已生成: {filename}")

if __name__ == "__main__":
    create_word()
    # create_pdf() # 如果你有 simhei.ttf 字体文件，可以取消注释这一行